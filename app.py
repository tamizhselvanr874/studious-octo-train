import os  
import re  
import streamlit as st  
from docx import Document  
from docx.shared import Pt  
import openai  
import io  
import fitz  # PyMuPDF  
  
# Azure OpenAI configuration  
openai.api_key = "783973291a7c4a74a1120133309860c0"  
openai.api_base = "https://theswedes.openai.azure.com/"  
openai.api_type = "azure"  
openai.api_version = "2024-05-01-preview"  
AZURE_DEPLOYMENT_NAME = "GPT-4-Omni"  
  
def extract_amendments(doc):  
    amendments = []  
    current_amendment = None  
    collecting = None  
    sections = {  
        "original": "",  
        "amended": "",  
        "source": "",  
        "reasoning": "",  
    }  
    supporting_argument = ""  
    found_supporting_argument = False  
    derivation_text = ""  
  
    for para in doc.paragraphs:  
        text = para.text.strip()  
  
        if text.startswith("Amendment"):  
            if current_amendment:  
                extract_source_and_reasoning(derivation_text, sections)  
                amendments.append({  
                    "title": current_amendment,  
                    **{key: value.strip() for key, value in sections.items()}  
                })  
            current_amendment = text  
            sections = {key: "" for key in sections}  
            collecting = None  
            derivation_text = ""  
        elif "Original Claim" in text:  
            collecting = "original"  
        elif "Proposed Amended Language" in text:  
            collecting = "amended"  
        elif "Derivation of Amendment" in text:  
            collecting = "derivation"  
        elif "Supporting Arguments" in text:  
            found_supporting_argument = True  
            collecting = None  
        elif collecting == "derivation":  
            derivation_text += text + "\n"  
        elif collecting:  
            sections[collecting] += text + "\n"  
        elif found_supporting_argument:  
            supporting_argument += text + "\n"  
  
    if current_amendment:  
        extract_source_and_reasoning(derivation_text, sections)  
        amendments.append({  
            "title": current_amendment,  
            **{key: value.strip() for key, value in sections.items()}  
        })  
  
    return amendments, supporting_argument.strip()  
  
def extract_source_and_reasoning(derivation_text, sections):  
    source_match = re.search(r"Source Reference:\s*(.*?)\s*(?=Reasoning:|$)", derivation_text, re.DOTALL)  
    reasoning_match = re.search(r"Reasoning:\s*(.*?)\s*($|^)", derivation_text, re.DOTALL)  
  
    if source_match:  
        sections["source"] = source_match.group(1).strip()  
    if reasoning_match:  
        sections["reasoning"] = reasoning_match.group(1).strip()  
  
def extract_text_from_pdf(pdf_data):  
    """Extract text from a PDF document using PyMuPDF."""  
    try:  
        pdf_document = fitz.open(stream=pdf_data, filetype="pdf")  
        text_by_page = {}  
  
        for page_number in range(len(pdf_document)):  
            page = pdf_document.load_page(page_number)  
            page_text = page.get_text("text")  
            text_by_page[page_number + 1] = page_text  
  
        pdf_document.close()  
        return text_by_page  
  
    except Exception as e:  
        st.error(f"An error occurred during PDF text extraction: {str(e)}")  
        return None  
  
def fetch_source_references(amendments, text_by_page):  
    all_text = "\n".join(text_by_page.values())  
    for amendment in amendments:  
        source_references = re.findall(r"\[(\d+)\]", amendment["source"])  
        fetched_content = []  
        for ref in source_references:  
            prompt = (  
                f"You are provided a document text with sections labeled with numbers in square brackets. "  
                f"Your task is to locate and extract the content following reference number [{ref}]. "  
                f"Please provide the paragraph or sentence directly associated with this reference:\n\n"  
                f"{all_text}"  
            )  
            try:  
                response = openai.ChatCompletion.create(  
                    engine=AZURE_DEPLOYMENT_NAME,  
                    messages=[  
                        {"role": "system", "content": "You are a knowledgeable assistant specialized in document analysis."},  
                        {"role": "user", "content": prompt}  
                    ],  
                    max_tokens=500  
                )  
                content = response.choices[0].message['content'].strip()  
                if content:  
                    fetched_content.append(f"Content from [{ref}]:\n{content}\n")  
                else:  
                    fetched_content.append(f"Content from [{ref}]: No content found.\n")  
            except Exception as e:  
                st.error(f"An error occurred while fetching source references: {str(e)}")  
                fetched_content.append(f"Content from [{ref}]: Error occurred.\n")  
        amendment["fetched_content"] = "\n".join(fetched_content)  
  
def compare_claims(original, amended):  
    prompt = (  
        f"Original Claim Language:\n{original}\n\n"  
        f"Proposed Amended Language:\n{amended}\n\n"  
        "Identify the additional parts in the Proposed Amended Language compared to the Original Claim Language."  
    )  
    try:  
        response = openai.ChatCompletion.create(  
            engine=AZURE_DEPLOYMENT_NAME,  
            messages=[  
                {"role": "system", "content": "You are a helpful assistant specialized in patent analysis."},  
                {"role": "user", "content": prompt}  
            ],  
            max_tokens=500  
        )  
        return response.choices[0].message['content'].strip()  
    except Exception as e:  
        st.error(f"An error occurred during claim comparison: {str(e)}")  
        return ""  
  
def generate_insights_for_amendment(amendment, supporting_argument):  
    prompt = (  
        f"Proposed Amendments and Arguments:\n"  
        f"Original Claim Language:\n{amendment['original']}\n\n"  
        f"Proposed Amended Language:\n{amendment['amended']}\n\n"  
        f"Derivation of Amendment:\n"  
        f"• *Source Reference*: {amendment['source']}\n"  
        f"• *Reasoning*: {amendment['reasoning']}\n\n"  
        "Please convert the provided \"Proposed Amendments and Arguments\" into the "  
        "\"Currently Amended\" claims format. Follow these instructions for each amendment:\n\n"  
        "1. *Format the Claim:*\n"  
        "   - Begin with the amendment number and label as \"(Currently Amended)\".\n"  
        "   - State the amended claim as a complete sentence, specifying the component "  
        "being amended and its new functionality or configuration.\n\n"  
        "2. *Derivation of Amendment:*\n"  
        "   - Include a subheading titled \"*Derivation of Amendment:*\".\n"  
        "   - Under this subheading, provide:\n"  
        "     - *Source Reference*: Indicate the paragraph or figure number from which "  
        "the amendment is derived.\n"  
        "     - *Reasoning*: Explain the purpose and advantage of the amendment, "  
        "particularly how it enhances specificity, novelty, or non-obviousness over prior art like Dwyer.\n\n"  
        "Example Output Format:\n\n"  
        "AMENDMENTS TO THE CLAIMS\n\n"  
        "1. (Currently Amended) [State the amended claim here.]\n"  
        "   - *Derivation of Amendment:*\n"  
        "     - *Source Reference*: [State the source reference here.]\n"  
        "     - *Reasoning*: [State the reasoning here.]\n\n"  
        "[Repeat for each amendment.]\n\n"  
        "Supporting Argument:\n{supporting_argument}\n\n"  
        "Provide a concise insight and detailed content for this amendment."  
    )  
    try:  
        response = openai.ChatCompletion.create(  
            engine=AZURE_DEPLOYMENT_NAME,  
            messages=[  
                {"role": "system", "content": "You are a helpful assistant specialized in patent insights."},  
                {"role": "user", "content": prompt}  
            ],  
            max_tokens=1000  
        )  
        return response.choices[0].message['content'].strip()  
    except Exception as e:  
        st.error(f"An error occurred during insight generation: {str(e)}")  
        return ""  
  
def format_content_in_patent_tone(extracted_content):  
    prompt = (  
        "Using a patent-like tone, combine the following extracted content into a single coherent passage. "  
        "Ensure the passage is detailed, structured, and flows logically, reflecting the comprehensive nature of patent disclosures. "  
        "Begin with a broad introduction, then integrate each cited element into a unified explanation, ending with a summary of the invention's utility.\n\n"  
        f"{extracted_content}"  
    )  
    try:  
        messages = [  
            {"role": "system", "content": "You are a helpful assistant skilled in technical writing."},  
            {"role": "user", "content": prompt}  
        ]  
        response = openai.ChatCompletion.create(  
            engine=AZURE_DEPLOYMENT_NAME,  
            messages=messages,  
            max_tokens=1500,  
            temperature=0.5  
        )  
        return response.choices[0].message['content'].strip()  
    except Exception as e:  
        st.error(f"An error occurred during content formatting: {str(e)}")  
        return ""  
  
def generate_reason_for_disagreement(amendment_contents, proposal_content):  
    prompt = (  
        "Please draft a response to an Office Action rejection regarding a patent application. "  
        "The response should respectfully disagree with the assertions made in the Office Action, "  
        "highlighting the following points:\n"  
        "- Identify how the cited references fail to disclose or suggest specific features introduced in the amended claims.\n"  
        "- Clearly outline how each proposed amendment introduces novel elements not present in the cited references.\n"  
        "- Emphasize the unique features and technical advantages introduced by each amendment.\n"  
        "- Assert the allowability of the amended independent claims over the cited art, and explain how dependent claims are patentable by virtue of their dependency on these allowable claims.\n"  
        "- Request reconsideration and withdrawal of the rejection.\n\n"  
        "The response should be structured with a respectful tone and include the following sections:\n"  
        "1. *Introduction*: Politely state the disagreement with the Office Action and provide a brief overview.\n"  
        "2. *Detailed Analysis*: For each amendment, explain:\n"  
        "   - The specific novel features and their derivation.\n"  
        "   - How the cited references do not suggest these features.\n"  
        "   - The technical advantages provided by these features.\n"  
        "3. *Conclusion*: Summarize the reasons for allowability and request the reconsideration and withdrawal of the rejection.\n\n"  
        "Amendment Contents:\n"  
        f"{amendment_contents}\n\n"  
        "Proposal Content:\n"  
        f"{proposal_content}\n\n"  
        "Please ensure the response is detailed, follows the specified format, and maintains a professional and respectful tone."  
    )  
    try:  
        messages = [  
            {"role": "system", "content": "You are a helpful assistant skilled in patent law."},  
            {"role": "user", "content": prompt}  
        ]  
        response = openai.ChatCompletion.create(  
            engine=AZURE_DEPLOYMENT_NAME,  
            messages=messages,  
            max_tokens=1500,  
            temperature=0.5  
        )  
        return response.choices[0].message['content'].strip()  
    except Exception as e:  
        st.error(f"An error occurred during reason generation: {str(e)}")  
        return ""  
  
def create_word_document(insights, reason_for_disagreement):  
    doc = Document()  
    for idx, insight in enumerate(insights, start=1):  
        doc.add_heading(f'Insights for Amendment {idx}', level=1)  
        p = doc.add_paragraph(insight)  
        p.style.font.size = Pt(12)  
  
    doc.add_heading('Reason for Disagreement', level=1)  
    p = doc.add_paragraph(reason_for_disagreement)  
    p.style.font.size = Pt(12)  
  
    return doc  
  
def main():  
    st.title("Patent Amendment Analyzer")  
  
    uploaded_word_file = st.file_uploader("Upload a Word Document", type="docx")  
    if uploaded_word_file is not None:  
        try:  
            doc = Document(uploaded_word_file)  
            amendments, supporting_argument = extract_amendments(doc)  
  
            # Indicate successful analysis  
            st.success("Word document analyzed successfully. Proceed to specify the number of PDF documents.")  
  
            # Ask for the number of PDFs to be uploaded  
            num_pdfs = st.number_input("Enter the number of PDF documents you want to upload:", min_value=1, step=1)  
  
            pdf_files = []  
            for i in range(num_pdfs):  
                uploaded_pdf_file = st.file_uploader(f"Upload PDF Document {i+1} (Application as Filed)", type="pdf", key=f"pdf_uploader_{i}")  
                if uploaded_pdf_file is not None:  
                    pdf_files.append(uploaded_pdf_file.read())  
  
            if len(pdf_files) == num_pdfs:  
                all_insights = []  
                all_amendment_contents = ""  
  
                for pdf_data in pdf_files:  
                    text_by_page = extract_text_from_pdf(pdf_data)  
  
                    if text_by_page:  
                        fetch_source_references(amendments, text_by_page)  
  
                        insights = []  
                        amendment_contents = ""  
                        for amendment in amendments:  
                            differences = compare_claims(amendment["original"], amendment["amended"])  
                            amendment["differences"] = differences  
  
                            insight = generate_insights_for_amendment(amendment, supporting_argument)  
                            insights.append(insight)  
                            amendment_contents += insight + "\n\n"  
  
                        all_insights.extend(insights)  
                        all_amendment_contents += amendment_contents  
  
                # Format the content in patent tone  
                proposal_content = format_content_in_patent_tone(all_amendment_contents)  
  
                # Generate reason for disagreement  
                reason_for_disagreement = generate_reason_for_disagreement(all_amendment_contents, proposal_content)  
  
                # Create and download Word document  
                doc = create_word_document(all_insights, reason_for_disagreement)  
                buffer = io.BytesIO()  
                doc.save(buffer)  
                buffer.seek(0)  
                st.download_button("Download Amendments", data=buffer, file_name="augmented_amendments.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")  
        except Exception as e:  
            st.error(f"An error occurred during document processing: {str(e)}")  
  
if __name__ == "__main__":  
    main()  